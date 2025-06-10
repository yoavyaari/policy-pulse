# src/app/apis/documents/__init__.py
import databutton as db
from fastapi import APIRouter, HTTPException, Depends, BackgroundTasks
from supabase.client import Client, create_client
from postgrest.exceptions import APIError
from typing import List, Optional, Dict, Any, Literal, get_args
import uuid
import traceback
from datetime import datetime, timezone
import json
from openai import OpenAI
import pypdf
import io
from io import BytesIO # Added BytesIO
from docx import Document # Added for DOCX processing
import asyncio
from pydantic import BaseModel, Field

# --- Model Definitions ---

# --- Model for PDF Processing Request ---
class ProcessPdfRequest(BaseModel):
    """Data needed to start processing a newly uploaded PDF."""
    storage_path: str = Field(..., description="The path to the PDF file in Supabase Storage.")
    user_id: str = Field(..., description="The ID of the user who uploaded the file.")
    file_name: str = Field(..., description="Original name of the uploaded file.")
    project_id: uuid.UUID = Field(..., description="The ID of the project this document belongs to.")

# --- Model for PDF Processing Response ---
class ProcessPdfResponse(BaseModel):
    """Response after initiating PDF processing."""
    success: bool
    message: str
    document_id: Optional[uuid.UUID] = None # ID of the created document record

class TopicDetail(BaseModel):
    name: Optional[str] = None
    sentiment: Optional[Literal["positive", "negative", "neutral"]] = None
    risks: Optional[List[str]] = Field(default_factory=list)
    regulation_needed: Optional[bool] = None

class DocumentAnalysis(BaseModel):
    submitter_name: Optional[str] = None
    response_date: Optional[str] = None # Consider validating date format if needed
    complexity_level: Optional[Literal["single sentence", "up to 2 paragraphs", "1-2 pages", "longer"]] = None
    depth_level: Optional[Literal["superficial", "moderate", "in-depth", "detailed"]] = None
    overall_sentiment: Optional[Literal["positive", "negative", "neutral"]] = None
    topics: Optional[List[TopicDetail]] = Field(default_factory=list)

class DocumentListItem(BaseModel):
    id: uuid.UUID
    file_name: str
    status: str
    complexity: Optional[Literal["single sentence", "up to 2 paragraphs", "1-2 pages", "longer"]] = Field(None)
    created_at: datetime
    project_id: uuid.UUID
    ai_analysis_error: Optional[str] = None
    processed_at: Optional[datetime] = None

    class Config:
        pass

class ListDocumentsResponse(BaseModel):
    documents: List[DocumentListItem]

class BasicReprocessResponse(BaseModel):
    success: bool
    message: str
    analysis_result: Optional[Dict[str, Any]] = None

class StepResult(BaseModel):
    step_id: uuid.UUID
    step_name: str
    prompt_template: Optional[str] = None
    result_data: Optional[Dict[str, Any]] = None
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    error_message: Optional[str] = None
    status: Literal["pending", "running", "completed", "error"] = "pending"

class DocumentDetailsResponse(BaseModel):
    id: uuid.UUID
    file_name: str
    status: str
    created_at: datetime
    processed_at: Optional[datetime] = None
    project_id: uuid.UUID
    storage_path: str
    user_id: Optional[str] = None
    analysis: Optional[DocumentAnalysis] = None
    custom_analysis_results: Optional[dict] = None # Renamed from custom_step_results
    ai_analysis_error: Optional[str] = None

class ReprocessRequest(BaseModel):
    reprocess_basic: bool = False
    reprocess_step_ids: Optional[List[uuid.UUID]] = None

class ReprocessResponse(BaseModel):
    success: bool
    message: str
    basic_analysis_status: Optional[str] = None
    step_reprocessing_status: Optional[Dict[uuid.UUID, str]] = None

class StepStats(BaseModel):
    step_id: uuid.UUID
    step_name: str
    completed_count: int
    pending_count: int
    error_count: int

class StepStatsResponse(BaseModel):
    stats: List[StepStats]

class IdentifyMissingComplexityResponse(BaseModel):
    document_ids: List[uuid.UUID]

class BulkBasicReprocessRequest(BaseModel):
    document_ids: Optional[List[uuid.UUID]] = None
    statuses: Optional[List[str]] = None  # Allow filtering by status
    project_id: Optional[uuid.UUID] = None # Allow scoping to a project

class BulkReprocessStartResponse(BaseModel):
    message: str
    task_count: int

class BulkFullReprocessRequest(BaseModel):
    document_ids: Optional[List[uuid.UUID]] = None

class FullReprocessResponse(BaseModel):
    success: bool
    message: str

class DeleteStepResultsResponse(BaseModel):
    success: bool
    message: str
    deleted_count: int

# --- End of Model Definitions ---

# --- Text Extraction Helper Functions (Sync) ---

def _extract_text_from_pdf_bytes_sync(pdf_bytes_io: BytesIO) -> Optional[str]:
    """Extracts text from PDF bytes using pypdf. Runs synchronously."""
    try:
        reader = pypdf.PdfReader(pdf_bytes_io)
        all_text_parts = [page.extract_text() for page in reader.pages if page.extract_text()] # Ensure text exists
        if not all_text_parts:
            return None # Or empty string, depending on desired behavior for empty PDFs
        return "\n\n".join(all_text_parts)
    except Exception as extraction_err:
        # Log the document_id if available, or a generic message
        print(f"[ERROR_SYNC_PDF_EXTRACTION] PDF parsing/extraction failed: {extraction_err}")
        # Optionally, re-raise or return None to indicate failure
        return None

def _extract_text_from_docx_bytes_sync(docx_bytes_io: BytesIO) -> Optional[str]:
    """Extracts text from DOCX bytes using python-docx. Runs synchronously."""
    try:
        document = Document(docx_bytes_io)
        all_text_parts = [para.text for para in document.paragraphs if para.text]
        if not all_text_parts:
            return None # Or empty string
        return "\n".join(all_text_parts)
    except Exception as extraction_err:
        print(f"[ERROR_SYNC_DOCX_EXTRACTION] DOCX parsing/extraction failed: {extraction_err}")
        return None

# Force reload comment 2025-05-03_22:08

# Force reload comment 2025-05-03_22:08

# --- Dependency Injection for Clients ---

def get_openai_client() -> OpenAI:
    """Initializes and returns an OpenAI client."""
    try:
        openai_api_key: str = db.secrets.get("OPENAI_API_KEY")
        if not openai_api_key:
            print("Error: OPENAI_API_KEY secret not found.")
            raise HTTPException(status_code=500, detail="Server configuration error: OpenAI API key missing.")
        return OpenAI(api_key=openai_api_key)
    except Exception as e:
        print(f"Error initializing OpenAI client: {e}")
        raise HTTPException(status_code=500, detail=f"Server configuration error: Could not initialize OpenAI client: {e}") from e

def get_supabase_client() -> Client:
    """Initializes and returns a Supabase client."""
    try:
        supabase_url: str = db.secrets.get("SUPABASE_URL")
        supabase_key: str = db.secrets.get("SUPABASE_SERVICE_ROLE_KEY")
        if not supabase_url or not supabase_key:
            print("Error: SUPABASE_URL or SUPABASE_SERVICE_ROLE_KEY secret not found.")
            raise HTTPException(status_code=500, detail="Server configuration error: Supabase secrets missing.")
        return create_client(supabase_url, supabase_key)
    except Exception as e:
        print(f"Error initializing Supabase client: {e}")
        raise HTTPException(status_code=500, detail=f"Server configuration error: Could not connect to Supabase: {e}") from e

router = APIRouter(prefix="/documents", tags=["documents"]) # Added prefix for clarity

# --- PDF Processing Background Task ---

async def _run_pdf_processing_task(
    supabase: Client,
    openai_client: OpenAI,
    document_id: uuid.UUID,
    storage_path: str,
    project_id: uuid.UUID, # Ensure project_id is passed
    user_id: str,
    file_name: str
):
    """Background task to download, analyze, and update the document."""
    print(f"[{document_id}] Background task started for {file_name} (Project: {project_id})")
    current_utc_time = datetime.now(timezone.utc).isoformat()
    analysis_result: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    extracted_text: Optional[str] = None # To store extracted text
    file_extension = file_name.split('.')[-1].lower()
    mime_type = "application/octet-stream" # Default
    if file_extension == "pdf":
        mime_type = "application/pdf"
    elif file_extension == "docx":
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    print(f"[{document_id}] Processing file {file_name} with detected mime_type: {mime_type}")

    try:
        # --- 1. Download File & Extract Text ---
        try:
            bucket_name = "pdf-documents" # TODO: Consider if bucket name needs to be dynamic or configurable if non-PDFs are stored elsewhere
            print(f"[{document_id}] Downloading from bucket '{bucket_name}', path '{storage_path}'...")
            storage_response = await asyncio.to_thread(
                 supabase.storage.from_(bucket_name).download, storage_path
            )
            file_bytes = storage_response # Renamed from pdf_bytes for generality
            if not file_bytes:
                raise ValueError("Downloaded file is empty or download failed.")
            print(f"[{document_id}] Downloaded {len(file_bytes)} bytes.")

            if mime_type == "application/pdf":
                def extract_text_pdf_sync(pdf_bytes_io):
                    try:
                        reader = pypdf.PdfReader(pdf_bytes_io)
                        all_text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
                        return "\n\n".join(all_text_parts)
                    except Exception as extraction_err:
                         print(f"[ERROR] PDF parsing failed during sync extraction for {document_id}: {extraction_err}")
                         raise ValueError(f"Failed to parse PDF content sync: {extraction_err}") from extraction_err
                extracted_text = await asyncio.to_thread(extract_text_pdf_sync, io.BytesIO(file_bytes))
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                def extract_text_docx_sync(docx_bytes_io):
                    try:
                        document = Document(docx_bytes_io)
                        all_text_parts = [para.text for para in document.paragraphs if para.text]
                        return "\n\n".join(all_text_parts)
                    except Exception as extraction_err:
                        print(f"[ERROR] DOCX parsing failed during sync extraction for {document_id}: {extraction_err}")
                        raise ValueError(f"Failed to parse DOCX content sync: {extraction_err}") from extraction_err
                extracted_text = await asyncio.to_thread(extract_text_docx_sync, io.BytesIO(file_bytes))
            else:
                raise ValueError(f"Unsupported mime_type for text extraction: {mime_type}")

            print(f"[{document_id}] Extracted {len(extracted_text)} characters.")
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from the file.")

        except Exception as file_err: # General file processing error
            print(f"[ERROR] Failed to download or extract text for doc {document_id}: {file_err}")
            error_message = f"Failed to download/extract file: {file_err}"
            raise # Re-raise to be caught by outer try-except

        # --- 2. Perform Basic Analysis (using the extracted text) ---
        analysis_result = await _perform_basic_analysis(
            supabase=supabase,
            openai_client=openai_client,
            document_id=document_id,
            storage_path=storage_path, # Still needed by helper in case of re-run without text
            extracted_text=extracted_text # Pass the extracted text
        )

        # --- 3. Update Document Record with analysis and status ---
        update_data = {
            "analysis": analysis_result,
            "status": "processed",
            "ai_analysis_error": None, # Clear previous errors
            "processed_at": current_utc_time,
            "extracted_text": extracted_text # Store extracted text
        }
        print(f"[{document_id}] Updating document with status 'processed' and analysis.")
        await asyncio.to_thread(
            supabase.table("documents")
            .update(update_data)
            .eq("id", str(document_id))
            .execute
        )
        print(f"[{document_id}] Background task completed successfully.")

    except Exception as task_err:
        # Catch errors from PDF download/extract or analysis helper
        if not error_message: # If not already set by PDF error
             error_message = f"Processing failed: {task_err}"
        print(f"[ERROR] Background task failed for {document_id}: {error_message}")
        traceback.print_exc()
        # Update document status to 'error'
        try:
            await asyncio.to_thread(
                supabase.table("documents")
                .update({
                    "status": "error",
                    "ai_analysis_error": error_message[:1000], # Truncate if needed
                    "processed_at": current_utc_time,
                    "extracted_text": extracted_text # Store even if analysis failed
                 })
                .eq("id", str(document_id))
                .execute
            )
            print(f"[{document_id}] Updated document status to 'error'.")
        except Exception as db_update_err:
            print(f"[ERROR] Failed to update document status to 'error' for {document_id}: {db_update_err}")


# --- Endpoint to Initiate PDF Processing ---
@router.post(
    "/process-pdf",
    response_model=ProcessPdfResponse,
    summary="Process Uploaded PDF",
    description="Creates a document record and starts background processing (download, text extraction, basic AI analysis) for a PDF uploaded to storage. Takes storage path, user ID, filename, and project ID as input."
)
async def process_pdf_endpoint(
    request: ProcessPdfRequest,
    background_tasks: BackgroundTasks,
    supabase: Client = Depends(get_supabase_client),
    openai_client: OpenAI = Depends(get_openai_client) # Inject clients
):
    """
    Receives PDF info, creates DB record, and triggers background analysis task.
    Includes project_id in the document creation.
    Determines mime_type based on file_name extension.
    """
    print(f"Received request to process file: {request.file_name} for project {request.project_id}")

    # Determine MIME type
    file_extension = request.file_name.split('.')[-1].lower()
    mime_type = "application/octet-stream" # Default
    if file_extension == "pdf":
        mime_type = "application/pdf"
    elif file_extension == "docx":
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    print(f"Determined MIME type: {mime_type} for file {request.file_name}")

    # 1. Create initial document record in Supabase
    try:
        # Include project_id and mime_type here
        insert_data = {
            "file_name": request.file_name,
            "storage_path": request.storage_path,
            "status": "uploaded", # Initial status
            "user_id": request.user_id,
            "project_id": str(request.project_id), # Store the project ID as string
            "mime_type": mime_type # Store the determined MIME type
        }
        print(f"Inserting document record: {insert_data}")
        response = await asyncio.to_thread(
            supabase.table("documents").insert(insert_data).execute
        )

        # Check response - Supabase client raises APIError on failure
        if not response.data:
            raise HTTPException(status_code=500, detail="Failed to create document record in database (no data returned).")

        document_data = response.data[0]
        document_id = document_data.get("id")
        if not document_id:
             raise HTTPException(status_code=500, detail="Failed to retrieve document ID after insertion.")

        print(f"Created document record with ID: {document_id}")

        # Convert document_id string to UUID object if necessary
        try:
            document_uuid = uuid.UUID(document_id)
        except ValueError:
             raise HTTPException(status_code=500, detail="Invalid document ID format received from database.")


        # 2. Add background task for processing
        background_tasks.add_task(
            _run_pdf_processing_task,
            supabase,
            openai_client,
            document_uuid, # Pass UUID object
            request.storage_path,
            request.project_id, # Pass project ID to task
            request.user_id,
            request.file_name
        )
        print(f"Added background task for document ID: {document_uuid}")

        # 3. Return success response
        return ProcessPdfResponse(
            success=True,
            message="Document upload accepted, processing started in background.",
            document_id=document_uuid
        )

    except APIError as api_error:
        print(f"Supabase API Error creating document record: {api_error}")
        raise HTTPException(status_code=500, detail=f"Database error: {api_error.message}")
    except HTTPException as http_exc:
        raise http_exc # Re-raise exceptions related to ID format etc.
    except Exception as e:
        print(f"Unexpected error processing PDF request: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An unexpected error occurred while initiating PDF processing.")


# --- Helper Function for Basic Analysis ---
# (Moved _perform_basic_analysis here, but it uses models defined in _models)
async def _perform_basic_analysis(
    supabase: Client,
    openai_client: OpenAI,
    document_id: uuid.UUID,
    storage_path: str,
    extracted_text: Optional[str] = None
) -> Dict[str, Any]:
    """
    Performs the initial LLM analysis using provided or extracted text.
    Returns the parsed JSON analysis result. Raises ValueError on failure.
    """
    print(f"[{document_id}] Starting basic analysis helper.")
    doc_content = extracted_text

    # Download/Extract only if text not provided (mainly for re-runs)
    if doc_content is None:
        file_bytes: Optional[bytes] = None
        bucket_name = "pdf-documents"  # Assuming this is the consistent bucket name
        try:
            print(f"[{document_id}] Attempting to download from Supabase Storage: bucket '{bucket_name}', path '{storage_path}'")
            # The Supabase client's download method is synchronous, so run in a thread
            # It typically returns bytes on success, or raises an APIError (like a 404 as an Object সংক্ষিপ্ত) if not found or access denied.
            # However, the exact error or return for "not found" can vary based on Supabase/storage-api versions.
            # For now, let's assume it might return None or empty bytes for a non-critical failure like not found, 
            # and raise an APIError for more critical issues.
            storage_response = await asyncio.to_thread(
                supabase.storage.from_(bucket_name).download, storage_path
            )
            file_bytes = storage_response

            if not file_bytes: # Check for empty bytes, common if download failed silently or file is empty/not found
                print(f"[{document_id}] Download from Supabase Storage returned empty or None. Path: {storage_path}")
                file_bytes = None # Ensure it's None to be caught by later check
            else:
                print(f"[{document_id}] Successfully downloaded {len(file_bytes)} bytes from Supabase Storage: {storage_path}.")
        except APIError as e_supabase_api:
            # Supabase client might raise APIError for various issues, including 404 (not found) or 403 (forbidden)
            # Check e_supabase_api.status or e_supabase_api.message for specifics if needed
            print(f"[{document_id}] Supabase API Error during download from path {storage_path}: Status {e_supabase_api.status if hasattr(e_supabase_api, 'status') else 'N/A'}, Message: {e_supabase_api.message}")
            file_bytes = None # Treat as download failure
        except Exception as e_download:
            print(f"[{document_id}] General error downloading from Supabase Storage path {storage_path}: {e_download}")
            traceback.print_exc() # Print full traceback for unexpected errors
            file_bytes = None # Treat as download failure

        if file_bytes:
            file_extension = storage_path.split('.')[-1].lower()
            try:
                if file_extension == "pdf":
                    print(f"[{document_id}] Attempting PDF text extraction for {storage_path}.")
                    doc_content = await asyncio.to_thread(_extract_text_from_pdf_bytes_sync, io.BytesIO(file_bytes))
                elif file_extension == "docx":
                    print(f"[{document_id}] Attempting DOCX text extraction for {storage_path}.")
                    doc_content = await asyncio.to_thread(_extract_text_from_docx_bytes_sync, io.BytesIO(file_bytes))
                else:
                    print(f"[{document_id}] Unsupported file type '{file_extension}' for text extraction from path {storage_path}.")
                    # doc_content remains None
                
                if doc_content:
                    print(f"[{document_id}] Successfully extracted text ({len(doc_content)} chars) from {storage_path}.")
                else:
                    print(f"[{document_id}] Text extraction yielded no content for {storage_path} (file type: {file_extension}).")
                    # doc_content is already None or will be set to None by the extraction helpers on failure/empty

            except Exception as e_extraction_call:
                # This catches errors in asyncio.to_thread or unexpected errors from sync helpers
                print(f"[{document_id}] Error during text extraction call for {storage_path}: {e_extraction_call}")
                doc_content = None # Ensure doc_content is None if extraction call fails

        # Check if doc_content is still None after attempted download/extraction
        if doc_content is None:
            error_message = f"Failed to retrieve or extract text content for document {document_id} from storage path {storage_path} during basic analysis."
            print(f"[{document_id}] {error_message} (Final check before raising ValueError)")
            raise ValueError(error_message)
        
        print(f"[{document_id}] Extracted text for analysis (length: {len(doc_content)} chars).")
    else: # doc_content was provided as an argument
        if doc_content is None: # Should not happen if called correctly, but a good safeguard
            error_message = f"Provided extracted_text was None for document {document_id}, which is unexpected when re-running basic analysis with provided text."
            print(f"[{document_id}] {error_message}")
            raise ValueError(error_message)
        print(f"[{document_id}] Using provided extracted text ({len(doc_content)} chars).")

    # Call LLM for Basic Analysis
    try:
        max_chars = 15000
        truncated_content = doc_content[:max_chars]
        # ... [Prompt definition as before - omitted for brevity] ...
        prompt = f"""Please analyze the following document content extracted from a policy response PDF. Provide the analysis ONLY as a valid JSON object containing the following keys:
- "submitter_name": (string) The name of the person or entity who submitted the response. If not found, use null.
- "response_date": (string) The date of the response in YYYY-MM-DD format. If not found, use null.
- "complexity_level": (string) Categorize the response complexity: "single sentence", "up to 2 paragraphs", "1-2 pages", "longer". If unsure, use null.
- "depth_level": (string) Assess the depth of analysis presented: "superficial", "moderate", "in-depth". If unsure, use null.
- "overall_sentiment": (string) The general sentiment towards the policy issue: "positive", "negative", "neutral". If unsure, use null.
- "topics": (array of objects) A list of general topics discussed. Each object should have:
    - "name": (string) Name of the topic.
    - "sentiment": (string) Sentiment for this specific topic ("positive", "negative", "neutral", or null).
    - "risks": (array of strings) List any risks presented for this topic. Use an empty array [] if none.
    - "regulation_needed": (boolean) Does the author suggest regulation is needed for this topic? Use true, false, or null if unclear.

Document Content:
```
{truncated_content}
```

Respond ONLY with the valid JSON object. Do not include explanations or markdown formatting.
"""

        completion = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI assistant performing initial analysis on policy documents. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        llm_response_content = completion.choices[0].message.content

        # Parse LLM Response
        try:
            analysis_result = json.loads(llm_response_content)
            # Simple validation: Check if it's a dict and has at least one expected key
            if not isinstance(analysis_result, dict) or "overall_sentiment" not in analysis_result:
                 raise ValueError("LLM response is not a valid JSON object or missing expected keys.")
            print(f"[{document_id}] Successfully parsed basic analysis from LLM.")
            # TODO: Validate with DocumentAnalysis model before returning?
            # validated_analysis = DocumentAnalysis(**analysis_result)
            # return validated_analysis.model_dump() # Return as dict
            return analysis_result
        except (json.JSONDecodeError, ValueError) as parse_err:
            print(f"[ERROR] Failed to parse/validate LLM JSON for doc {document_id}: {parse_err}. Response: {llm_response_content}")
            raise ValueError(f"Failed to parse/validate LLM analysis response: {parse_err}") from parse_err

    except Exception as llm_err:
        print(f"[ERROR] Failed during LLM basic analysis call for doc {document_id}: {llm_err}")
        raise ValueError(f"LLM analysis failed: {llm_err}") from llm_err


# --- Endpoint to List Documents ---
@router.get(
    "/", # Use root path under /documents prefix
    response_model=ListDocumentsResponse,
    summary="List Documents by Project",
    description="Fetches a list of documents, optionally filtered by the currently selected project ID."
)
async def list_documents( # Changed to async
    project_id: uuid.UUID, # Make project_id mandatory path or query parameter? Query for now.
    supabase: Client = Depends(get_supabase_client)
) -> ListDocumentsResponse:
    """Lists documents filtered by project_id."""
    print(f"Received request to list documents for project_id: {project_id}")
    try:
        # Fetch necessary fields including complexity directly if stored at top level
        # If complexity is nested in 'analysis', adjust the select: 'analysis->>complexity_level'
        query = supabase.table("documents").select(
            "id, file_name, status, created_at, project_id, ai_analysis_error, analysis, processed_at" # Fetch analysis blob and processed_at
        ).eq("project_id", str(project_id)).order("created_at", desc=True)

        # Run query using asyncio.to_thread
        response = await asyncio.to_thread(query.execute)

        if response.data is not None:
            documents_data = response.data
            print(f"Found {len(documents_data)} documents for project {project_id}.")
            
            # Process data to extract complexity and validate
            processed_docs = []
            for doc in documents_data:
                analysis_data = doc.get("analysis") or {}
                complexity = analysis_data.get("complexity_level")
                # Ensure complexity matches the Literal options or is None
                valid_complexities = get_args(DocumentListItem.__annotations__['complexity'].__args__[0])
                if complexity not in valid_complexities:
                    complexity = None # Default to None if invalid

                processed_docs.append(DocumentListItem(
                    id=doc['id'],
                    file_name=doc['file_name'],
                    status=doc['status'],
                    created_at=doc['created_at'],
                    project_id=doc['project_id'],
                    ai_analysis_error=doc.get('ai_analysis_error'),
                    complexity=complexity, # Assign extracted/validated complexity
                    processed_at=doc.get('processed_at')
                ))
                
            return ListDocumentsResponse(documents=processed_docs)
        else:
            print(f"No documents found or Supabase response data is None for project {project_id}.")
            return ListDocumentsResponse(documents=[])

    except APIError as api_error:
        print(f"Supabase API Error listing documents for project {project_id}: {api_error}")
        raise HTTPException(status_code=500, detail=f"Database error listing documents: {api_error.message}")
    except Exception as e:
        print(f"Unexpected error listing documents for project {project_id}: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An unexpected error occurred while listing documents.")


# --- Endpoint for Basic Reprocessing ---
@router.post(
    "/{document_id}/reprocess-basic",
    response_model=BasicReprocessResponse,
    summary="Reprocess Basic Analysis",
    description="Downloads a PDF, re-runs the standard initial analysis, and updates the 'analysis' field for the specified document ID."
)
async def reprocess_basic_analysis_endpoint(
    document_id: uuid.UUID,
    supabase: Client = Depends(get_supabase_client),
    openai_client: OpenAI = Depends(get_openai_client)
):
    """
    Endpoint to trigger basic reprocessing for a single document.
    """
    print(f"Received request to reprocess basic analysis for document ID: {document_id}")
    # 1. Fetch document storage_path
    storage_path = await _get_doc_storage_path(document_id, supabase)

    # 2. Perform basic analysis using the helper
    analysis_result: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    current_utc_time = datetime.now(timezone.utc).isoformat()
    try:
        analysis_result = await _perform_basic_analysis(supabase, openai_client, document_id, storage_path)
    except ValueError as analysis_err: # Catch specific errors from helper
        error_message = str(analysis_err)
        print(f"[ERROR] Basic analysis failed for {document_id}: {error_message}")
        await _update_doc_status(document_id, supabase, "error", error_message, current_utc_time)
        return BasicReprocessResponse(success=False, message=f"Analysis failed: {error_message}")
    except Exception as e:
        error_message = "An unexpected error occurred during analysis."
        print(f"[ERROR] Unexpected error during basic analysis call for {document_id}: {e}")
        traceback.print_exc()
        await _update_doc_status(document_id, supabase, "error", error_message, current_utc_time)
        return BasicReprocessResponse(success=False, message=error_message)

    # 3. Update document in Supabase if analysis succeeded
    if analysis_result:
        try:
            await _update_doc_analysis(document_id, supabase, analysis_result, current_utc_time)
            print(f"Successfully updated basic analysis for document {document_id}")
            return BasicReprocessResponse(success=True, message="Basic analysis reprocessed successfully.", analysis_result=analysis_result)
        except (APIError, HTTPException) as db_update_error:
            error_message = f"Database update failed: {getattr(db_update_error, 'detail', getattr(db_update_error, 'message', str(db_update_error)))}"
            print(f"[ERROR] Failed to update Supabase after basic reprocess for doc {document_id}: {error_message}")
            await _update_doc_status(document_id, supabase, "error", error_message) # Update status to error
            return BasicReprocessResponse(success=False, message=error_message, analysis_result=analysis_result) # Include analysis if available
        except Exception as db_e:
            error_message = "Analysis succeeded but an unexpected error occurred during database update."
            print(f"[ERROR] Unexpected error updating Supabase for doc {document_id}: {db_e}")
            traceback.print_exc()
            # Don't update status here as analysis itself succeeded
            return BasicReprocessResponse(success=False, message=error_message, analysis_result=analysis_result)

    # Fallback if analysis_result was somehow None without error
    return BasicReprocessResponse(success=False, message="Analysis did not complete successfully.")


# --- Bulk Reprocessing Endpoints (Simplified Stubs) ---
# (Keep existing bulk endpoints, ensure they use helpers correctly)

@router.post(
    "/bulk-reprocess-basic",
    response_model=BulkReprocessStartResponse,
    tags=["bulk"],
    summary="Start Bulk Reprocess Basic Analysis",
    description="Starts basic analysis reprocessing for specified documents, or all documents in a project, optionally filtered by status."
)
async def trigger_bulk_basic_reprocessing(
    request: BulkBasicReprocessRequest,
    background_tasks: BackgroundTasks,
    supabase_client: Client = Depends(get_supabase_client),
    openai_client: OpenAI = Depends(get_openai_client)  # Added openai_client
) -> BulkReprocessStartResponse:
    """Triggers a background task for bulk basic reprocessing of documents."""
    print(f"Received bulk basic reprocess request: {request}")

    if not request.document_ids and not request.project_id:
        # If project_id is also not provided, then it's an invalid request.
        # The UI should always provide project_id for a general 'reprocess all in project'.
        raise HTTPException(
            status_code=400,
            detail="Either document_ids or project_id must be provided for bulk reprocessing."
        )

    try:
        eligible_doc_ids = await _get_eligible_doc_ids(
            supabase_client=supabase_client,
            document_ids=request.document_ids,
            statuses=request.statuses, # Pass statuses if provided
            project_id=request.project_id # Pass project_id if provided
        )

        if not eligible_doc_ids:
            return BulkReprocessStartResponse(message="No documents found matching the criteria for reprocessing.", task_count=0)

        print(f"Queuing bulk basic reprocessing for {len(eligible_doc_ids)} documents.")
        background_tasks.add_task(
            _run_bulk_basic_reprocessing_task,
            eligible_doc_ids,
            supabase_client, # Pass Supabase client
            openai_client    # Pass OpenAI client
        )
        return BulkReprocessStartResponse(message=f"Bulk basic reprocessing initiated for {len(eligible_doc_ids)} documents.", task_count=len(eligible_doc_ids))
    except HTTPException as http_exc:
        raise http_exc # Re-raise HTTP exceptions
    except Exception as e:
        print(f"Error during bulk basic reprocessing trigger: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")


# --- Helper DB Functions ---

async def _get_doc_storage_path(doc_id: uuid.UUID, supabase: Client) -> str:
    """Fetches storage_path for a document ID. Raises 404 if not found."""
    try:
        response = await asyncio.to_thread(
             supabase.table("documents").select("storage_path").eq("id", str(doc_id)).limit(1).execute
        )
        if not response.data or not response.data[0].get("storage_path"):
            raise HTTPException(status_code=404, detail=f"Document or storage path not found for ID {doc_id}.")
        return response.data[0]["storage_path"]
    except APIError as api_err:
        print(f"[DB Helper ERROR] Fetching path for {doc_id}: {api_err}")
        raise HTTPException(status_code=500, detail=f"Database error fetching document path: {api_err.message}")
    except Exception as e:
        print(f"[DB Helper ERROR] Unexpected error fetching path for {doc_id}: {e}")
        raise HTTPException(status_code=500, detail="Unexpected error fetching document path.")

async def _update_doc_status(doc_id: uuid.UUID, supabase: Client, status: str, error_message: Optional[str] = None, timestamp: Optional[str] = None):
    """Updates document status and error message."""
    update_payload = {"status": status, "ai_analysis_error": error_message[:1000] if error_message else None}
    if timestamp:
        update_payload["processed_at"] = timestamp
    try:
        await asyncio.to_thread(
            supabase.table("documents").update(update_payload).eq("id", str(doc_id)).execute
        )
        print(f"[{doc_id}] Updated status to '{status}'. Error: {error_message if error_message else 'None'}")
    except Exception as db_err:
        print(f"[DB Helper ERROR] Failed to update status to '{status}' for {doc_id}: {db_err}")
        # Don't raise here, just log the failure

async def _update_doc_analysis(doc_id: uuid.UUID, supabase: Client, analysis_result: Dict[str, Any], timestamp: str):
    """Updates document analysis, status, and clears error."""
    update_data = {
        "analysis": analysis_result,
        "status": "processed",
        "ai_analysis_error": None,
        "processed_at": timestamp
    }
    try:
        update_response = await asyncio.to_thread(
            supabase.table("documents").update(update_data).eq("id", str(doc_id)).execute
        )
        # Optional: Check update_response.data for confirmation if needed
        if not update_response.data:
             print(f"[WARN] Supabase update analysis for {doc_id} returned no data.")
             # Consider raising an error if confirmation is critical
             # raise HTTPException(status_code=404, detail=f"Document {doc_id} not found for analysis update.")
    except APIError as db_update_error:
        print(f"[DB Helper ERROR] Failed to update analysis for {doc_id}: {db_update_error}")
        raise # Re-raise APIError to be handled by calling endpoint
    except Exception as e:
        print(f"[DB Helper ERROR] Unexpected error updating analysis for {doc_id}: {e}")
        raise # Re-raise unexpected errors

async def _get_eligible_doc_ids(
    supabase_client: Client, 
    document_ids: Optional[List[uuid.UUID]] = None, 
    statuses: Optional[List[str]] = None,
    project_id: Optional[uuid.UUID] = None
) -> List[uuid.UUID]:
    """Fetches document IDs eligible for reprocessing based on IDs, statuses, or project."""
    if not project_id and not document_ids:
        print("[DB Helper WARN] _get_eligible_doc_ids called without project_id or document_ids. Returning empty list.")
        return [] # Avoid fetching all documents by mistake

    query = supabase_client.table("documents").select("id, project_id") # Select project_id for verification if needed

    if project_id:
        print(f"[DB Helper] Filtering eligible docs by project_id: {project_id}")
        query = query.eq("project_id", str(project_id))
        # If specific document_ids are also given, they should be within this project
        if document_ids:
            str_document_ids = [str(doc_id) for doc_id in document_ids]
            query = query.in_("id", str_document_ids)
    elif document_ids: # Only use if project_id is not specified
        print(f"[DB Helper] Filtering eligible docs by specific document_ids: {document_ids}")
        str_document_ids = [str(doc_id) for doc_id in document_ids]
        query = query.in_("id", str_document_ids)
    # If neither project_id nor document_ids are provided, we've returned empty already.

    if statuses:
        print(f"[DB Helper] Additionally filtering by statuses: {statuses}")
        query = query.in_("status", statuses)
    
    try:
        response = await asyncio.to_thread(query.execute)
        if response.data:
            # Return only the IDs
            eligible_ids = [item['id'] for item in response.data]
            print(f"[DB Helper] Found {len(eligible_ids)} eligible documents.")
            return eligible_ids
        print("[DB Helper] No eligible documents found for the given criteria.")
        return []
    except APIError as api_err:
        print(f"[DB Helper ERROR] APIError fetching eligible document IDs: {api_err}")
        raise HTTPException(status_code=500, detail=f"Database error: {api_err.message}")
    except Exception as e:
        print(f"[DB Helper ERROR] Unexpected error fetching eligible document IDs: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error fetching documents: {str(e)}")


# --- Bulk Processing Helper Tasks ---
async def _run_bulk_basic_reprocessing_task(document_ids: List[uuid.UUID], supabase_client: Client, openai_client: OpenAI):
    """The actual background task that performs basic reprocessing for each document."""
    print(f"BG TASK: Starting basic reprocessing for {len(document_ids)} documents.")
    processed_count = 0
    error_count = 0

    for doc_id_str in document_ids: # Assuming document_ids are already strings or UUIDs
        doc_id = uuid.UUID(str(doc_id_str)) # Ensure it's a UUID object
        print(f"BG TASK: Reprocessing document ID: {doc_id}")
        try:
            # 1. Get document's file_name and project_id to construct storage_path
            doc_details_response = await asyncio.to_thread(
                supabase_client.table("documents").select("file_name, project_id, storage_path") # Fetch storage_path too if available
                .eq("id", str(doc_id)).maybe_single().execute
            )

            if not doc_details_response.data:
                print(f"BG TASK ERROR: Document ID {doc_id} not found in database.")
                await _update_doc_status(doc_id, supabase_client, "error", error_message=f"Document not found for reprocessing.")
                error_count += 1
                continue
            
            doc_data = doc_details_response.data
            file_name = doc_data.get("file_name")
            project_uuid = doc_data.get("project_id")
            # Prefer direct storage_path if already stored and correct, otherwise construct it.
            # Based on process_pdf_endpoint, storage_path in DB is the full path for Supabase storage
            # e.g., "public/717c9383-6782-4818-b9c1-c2b46ddb1abc/projects/PROJECT_ID_HERE/uploads/file.pdf"
            # The `storage_path` in `process_pdf_endpoint` is `request.storage_path` which comes from the client.
            # The client constructs it as: `const filePath = `public/${projectId}/projects/${currentProjectUuid}/uploads/${file.name}`;`
            # The `storage_path` field in the `documents` table should be this direct path.
            
            storage_key = doc_data.get("storage_path")
            if not storage_key:
                 # Fallback: construct if not in DB, though it should be.
                 # This logic assumes files are stored under `projects/{project_id}/uploads/{file_name}` in the bucket.
                 # And that the bucket itself is not part of the key stored/retrieved via db.storage.binary.
                 # db.storage.binary.get uses a key that does NOT include the bucket name.
                 # e.g. `projects/PROJECT_ID_HERE/uploads/file.pdf`
                 # Let's re-verify where PdfUploader puts files.
                 # `PdfUploader.tsx` has `const databuttonStoragePath = `projects/${currentProjectId}/uploads/${file.name}`;`
                 # `db.storage.binary.put(databuttonStoragePath, ...)` -> so `storage_key` for db.storage is this path.
                 # The `storage_path` in DB from `process_pdf_endpoint` was `request.storage_path` which is the full Supabase Storage path.
                 # This is confusing. Let's clarify storage: 
                 #   - `db.storage` uses keys like `my-folder/my-file.txt`.
                 #   - Supabase Storage uses paths like `bucket-name/my-folder/my-file.txt`.
                 # `_run_pdf_processing_task` uses `supabase.storage.from_(bucket_name).download(storage_path)`
                 # where `storage_path` is the Supabase Storage path (e.g. `public/...`).
                 # `_perform_basic_analysis` was designed to take `storage_path` (Supabase one) to re-download if text not provided.
                 # However, `_perform_basic_analysis` should ideally work with bytes directly if text is not provided.
                 # Let's fetch the bytes here directly using the db.storage key, assuming file_name + project_id is enough to build it.
                
                if not file_name or not project_uuid:
                    print(f"BG TASK ERROR: Document {doc_id} missing file_name or project_id for storage key construction.")
                    await _update_doc_status(doc_id, supabase_client, "error", error_message="Missing info for storage key.")
                    error_count += 1
                    continue
                # THIS IS THE KEY FOR db.storage, NOT Supabase Storage.
                # Ensure this matches what PdfUploader uses with db.storage.binary.put if that was the source.
                # If source is Supabase Storage via frontend upload, then we need the Supabase path.
                # The `process_pdf_endpoint` receives a `storage_path` (Supabase one) and stores it.
                # So, we should use `doc_data.get("storage_path")` and download from Supabase storage.
                storage_path_supabase = doc_data.get("storage_path") # This IS the Supabase Storage path from upload.
                if not storage_path_supabase:
                    print(f"BG TASK ERROR: Document {doc_id} missing storage_path (Supabase path).")
                    await _update_doc_status(doc_id, supabase_client, "error", error_message="Missing Supabase storage path.")
                    error_count += 1
                    continue
                
                bucket_name = "pdf-documents" # As used in _run_pdf_processing_task
                print(f"BG TASK: Downloading for {doc_id} from Supabase Storage: bucket '{bucket_name}', path '{storage_path_supabase}'")
                pdf_bytes_response = await asyncio.to_thread(
                    supabase_client.storage.from_(bucket_name).download, storage_path_supabase
                )
                pdf_bytes = pdf_bytes_response
            else: # storage_key is populated, assume it's the Supabase Storage path
                bucket_name = "pdf-documents"
                print(f"BG TASK: Downloading for {doc_id} from Supabase Storage: bucket '{bucket_name}', path '{storage_key}'")
                pdf_bytes_response = await asyncio.to_thread(
                    supabase_client.storage.from_(bucket_name).download, storage_key
                )
                pdf_bytes = pdf_bytes_response

            if not pdf_bytes:
                print(f"BG TASK ERROR: PDF content not found for document {doc_id} (path: {storage_key or storage_path_supabase}).")
                await _update_doc_status(doc_id, supabase_client, "error", error_message=f"PDF content not found.")
                error_count += 1
                continue
            
            # Extract text from PDF bytes
            extracted_text: Optional[str] = None
            try:
                def extract_text_sync(pdf_bytes_io_local):
                    reader_local = pypdf.PdfReader(pdf_bytes_io_local)
                    all_text_parts_local = [page.extract_text() for page in reader_local.pages if page.extract_text()]
                    return "\n\n".join(all_text_parts_local)
                extracted_text = await asyncio.to_thread(extract_text_sync, io.BytesIO(pdf_bytes))
                if not extracted_text or not extracted_text.strip():
                    raise ValueError("No text could be extracted from the PDF for reprocessing.")
                print(f"BG TASK: Extracted {len(extracted_text)} chars for doc {doc_id}.")
            except Exception as text_extract_err:
                print(f"BG TASK ERROR: Failed to extract text for doc {doc_id}: {text_extract_err}")
                await _update_doc_status(doc_id, supabase_client, "error", error_message=f"Text extraction failed: {text_extract_err}")
                error_count += 1
                continue

            # 2. Perform basic analysis and get the result
            analysis_result = await _perform_basic_analysis(
                supabase=supabase_client, 
                openai_client=openai_client, 
                document_id=doc_id, 
                storage_path=storage_key or storage_path_supabase, # Pass the Supabase storage path used for download
                extracted_text=extracted_text
            )
            
            # 3. Update the document with the new analysis and processed_at timestamp
            current_utc_time_str = datetime.now(timezone.utc).isoformat()
            await _update_doc_analysis(doc_id, supabase_client, analysis_result, current_utc_time_str)
            
            print(f"BG TASK: Successfully reprocessed and updated document ID: {doc_id}")
            processed_count += 1

        except Exception as e:
            print(f"BG TASK ERROR: Error reprocessing document {doc_id}: {e}")
            traceback.print_exc()
            await _update_doc_status(doc_id, supabase_client, "error", error_message=f"Reprocessing failed: {str(e)}")
            error_count += 1
            
    print(f"BG TASK: Finished basic reprocessing for {len(document_ids)} documents. Processed: {processed_count}, Errors: {error_count}.")

# Add other endpoints like get_document_details, reprocess_full, etc. as needed
# Ensure they also use the models imported from ._models

@router.get("/{document_id}", response_model=DocumentDetailsResponse, summary="Get Document Details")
async def get_document_details(
    document_id: uuid.UUID,
    project_id: uuid.UUID, # Added project_id query param
    supabase: Client = Depends(get_supabase_client)
):
    """
    Fetches details for a specific document, ensuring it belongs to the specified project.
    Retrieves analysis results directly from the 'documents.analysis' JSON field.
    """
    print(f"Fetching details for document {document_id} in project {project_id}")
    try:
        # 1. Fetch main document details, including the 'analysis' field
        doc_resp = await asyncio.to_thread(
            supabase.table("documents")
            .select("*") # Fetch all columns, including 'analysis'
            .eq("id", str(document_id))
            .eq("project_id", str(project_id)) # Filter by project_id
            .maybe_single()
            .execute
        )
        if not doc_resp.data:
            raise HTTPException(status_code=404, detail="Document not found in the specified project")
        
        document_data = doc_resp.data

        # 2. Validate analysis data structure directly from fetched data
        analysis_data = document_data.get("analysis")
        validated_analysis: Optional[DocumentAnalysis] = None
        if isinstance(analysis_data, dict):
            try:
                validated_analysis = DocumentAnalysis(**analysis_data)
            except Exception as validation_error:
                print(f"[WARN] Document {document_id}: Analysis data failed validation: {validation_error}")
                # Keep validated_analysis as None if validation fails
        elif analysis_data is not None:
             print(f"[WARN] Document {document_id}: Analysis data exists but is not a dictionary: {type(analysis_data)}")

        # 3. Prepare response
        #    custom_analysis_results is already included in document_data from the SELECT *

        # 4. Return combined response
        return DocumentDetailsResponse(
            **document_data, # Unpack main document data
            # analysis=validated_analysis, # REMOVED - Already included in **document_data
            # custom_analysis_results is already included via **document_data
        )
            
    except APIError as e:
        # Specific database API errors
        print(f"[ERROR] Database API error fetching details for doc {document_id} / project {project_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Database error: {e.message}")
    except HTTPException as http_exc:
        # Re-raise specific HTTP exceptions (like 404)
        raise http_exc 
    except Exception as e:
        # Catch-all for other unexpected errors
        print(f"[ERROR] Unexpected error fetching details for doc {document_id} / project {project_id}: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An unexpected error occurred while fetching document details.")


# Placeholder for utility function to get Literal args (if needed, e.g., for validation)
from typing import get_args

